
import io
from datetime import datetime
from pathlib import Path

from .config import LOGO_PATH


def _logo_bytes() -> bytes | None:
    if LOGO_PATH.exists():
        return LOGO_PATH.read_bytes()
    return None


def exportar_docx(dados: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo = _logo_bytes()
    if logo:
        run = hp.add_run()
        run.add_picture(io.BytesIO(logo), width=Inches(0.6))
        run.add_text("  ")
    run = hp.add_run("Polícia Judiciária — UNCT | OCR + Tradução")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x65, 0x6d, 0x76)

    titulo = doc.add_heading("Relatório de Tradução", level=1)
    titulo.runs[0].font.color.rgb = RGBColor(0x1b, 0x1f, 0x23)

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Light Grid Accent 1"
    meta_data = [
        ("Data", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Ficheiro", dados.get("ficheiro", "N/A")),
        ("Língua", dados.get("lingua", "N/A")),
        ("Modelo", dados.get("modelo", "N/A")),
        ("Caracteres OCR", str(len(dados.get("texto_ocr", "")))),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_paragraph()

    paragrafos_orig = dados.get("paragrafos_orig", [])
    paragrafos_pt = dados.get("paragrafos_pt", [])
    paragrafos_en = dados.get("paragrafos_en", [])

    if paragrafos_orig and paragrafos_pt:
        doc.add_heading("Tradução por parágrafos", level=2)
        cols = 3 if paragrafos_en else 2
        headers = ["Original", "Português (PT-PT)"]
        if paragrafos_en:
            headers.insert(1, "Inglês (intermédio)")

        table = doc.add_table(rows=1, cols=cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

        n = min(len(paragrafos_orig), len(paragrafos_pt))
        for i in range(n):
            row = table.add_row()
            row.cells[0].text = paragrafos_orig[i]
            if paragrafos_en:
                row.cells[1].text = paragrafos_en[i] if i < len(paragrafos_en) else ""
                row.cells[2].text = paragrafos_pt[i]
            else:
                row.cells[1].text = paragrafos_pt[i]
    else:
        if dados.get("texto_ocr"):
            doc.add_heading("Texto OCR", level=2)
            doc.add_paragraph(dados["texto_ocr"])
        if dados.get("texto_en"):
            doc.add_heading("Inglês (intermédio)", level=2)
            doc.add_paragraph(dados["texto_en"])
        if dados.get("texto_pt"):
            doc.add_heading("Português (PT-PT)", level=2)
            doc.add_paragraph(dados["texto_pt"])

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Pipeline OCR + Tradução — PJ/UNCT — Documento gerado automaticamente")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x8b, 0x94, 0x9e)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def exportar_excel(dados: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Tradução"

    headers = ["#", "Original", "Inglês", "Português (PT-PT)"]
    header_fill = PatternFill(start_color="1b1f23", end_color="1b1f23", fill_type="solid")
    header_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    paragrafos_orig = dados.get("paragrafos_orig", [])
    paragrafos_en = dados.get("paragrafos_en", [])
    paragrafos_pt = dados.get("paragrafos_pt", [])

    n = max(len(paragrafos_orig), len(paragrafos_pt))
    for i in range(n):
        row = i + 2
        ws.cell(row=row, column=1, value=i + 1)
        ws.cell(row=row, column=2, value=paragrafos_orig[i] if i < len(paragrafos_orig) else "")
        ws.cell(row=row, column=3, value=paragrafos_en[i] if i < len(paragrafos_en) else "")
        ws.cell(row=row, column=4, value=paragrafos_pt[i] if i < len(paragrafos_pt) else "")

    confianca = dados.get("confianca_ocr", [])
    if confianca:
        ws2 = wb.create_sheet("Confiança OCR")
        headers2 = ["Palavra", "Confiança (%)", "Bloco", "Linha"]
        for col, h in enumerate(headers2, 1):
            cell = ws2.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
        for i, item in enumerate(confianca):
            row = i + 2
            ws2.cell(row=row, column=1, value=item.get("palavra", ""))
            ws2.cell(row=row, column=2, value=item.get("confianca", 0))
            ws2.cell(row=row, column=3, value=item.get("bloco", 0))
            ws2.cell(row=row, column=4, value=item.get("linha", 0))

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 50

    ws_meta = wb.create_sheet("Metadados")
    meta = [
        ("Data", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Ficheiro", dados.get("ficheiro", "N/A")),
        ("Língua", dados.get("lingua", "N/A")),
        ("Modelo", dados.get("modelo", "N/A")),
    ]
    for i, (k, v) in enumerate(meta):
        ws_meta.cell(row=i + 1, column=1, value=k).font = Font(bold=True)
        ws_meta.cell(row=i + 1, column=2, value=v)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def exportar_pdf(dados: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage)
    except ImportError:
        raise RuntimeError("reportlab não instalado. python -m pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="PTResult",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="Meta",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.grey,
    ))

    story = []

    story.append(Paragraph("Relatório de Tradução — PJ/UNCT", styles["Title"]))
    story.append(Spacer(1, 12))

    meta_text = (
        f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Ficheiro: {dados.get('ficheiro', 'N/A')} | "
        f"Língua: {dados.get('lingua', 'N/A')} | "
        f"Modelo: {dados.get('modelo', 'N/A')}"
    )
    story.append(Paragraph(meta_text, styles["Meta"]))
    story.append(Spacer(1, 18))

    paragrafos_orig = dados.get("paragrafos_orig", [])
    paragrafos_pt = dados.get("paragrafos_pt", [])

    if paragrafos_orig and paragrafos_pt:
        table_data = [["#", "Original", "Português (PT-PT)"]]
        n = min(len(paragrafos_orig), len(paragrafos_pt))
        for i in range(n):
            table_data.append([
                str(i + 1),
                Paragraph(paragrafos_orig[i][:500], styles["PTResult"]),
                Paragraph(paragrafos_pt[i][:500], styles["PTResult"]),
            ])

        col_widths = [1*cm, 7.5*cm, 7.5*cm]
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1b1f23")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d7de")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fa")]),
        ]))
        story.append(table)
    else:
        if dados.get("texto_pt"):
            story.append(Paragraph("Resultado (PT-PT):", styles["Heading2"]))
            story.append(Paragraph(dados["texto_pt"][:3000], styles["PTResult"]))

    story.append(Spacer(1, 24))
    story.append(Paragraph(
        "Pipeline OCR + Tradução — PJ/UNCT — Documento gerado automaticamente",
        styles["Meta"]))

    doc.build(story)
    return buf.getvalue()
